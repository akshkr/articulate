import os
import docx
import time


class DocSavar:
	
	def __init__(self):
		self.doc = docx.Document()
	
	def save_to_doc(self, input_json, pdf = False):
		try:
			for key, value in input_json.items():
				self.doc.add_heading(key, 0)
				self.doc.add_paragraph(value)
			document_name = str(time.time())
			self.doc.save(document_name + '.docx')
			print("The file is saved in %s as %s" %(os.getcwd(), document_name+'.docx'))
			if pdf:
				return document_name+'.docx'
		except Exception as ex:
			print("Error encountered while saving docx")
			print(ex)
