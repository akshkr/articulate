"""
Article Maker: Based upon the user article topic article maker creates a rough article by gathering the information from wikipeida.
"""
#packages required
import sys
import time
import wikipedia
import docx
import queue
import requests
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtCore import *
# creating instance of word document.
doc = docx.Document()
# Base class
class App(QMainWindow):

    def start(self):
	    #hello
        print("")
    # constructor of the App class. 
    def __init__(self):
        super().__init__()
        self.title = 'Articulate - tecbeast.com'
        self.left = 50
        self.top = 100
        self.width = 400
        self.height = 300
        self.initUI()
    # Qt Prompt window customization.
    def initUI(self):
        self.setWindowTitle(self.title) # set the title
        self.setGeometry(self.left, self.top, self.width, self.height) # set the dimensions of the prompt screen
        lbl1 = QLabel(self) # object creation of the Qlabel
        lbl1.setText("Please enter the topics you want to print articles on") # set Text of the label 
        lbl1.setGeometry(20,20,400,40) # set the position of the label on the prompt screen.

        lbl2 = QLabel(self)
        lbl2.setText('Topics should be separeted by comma ","')
        lbl2.setGeometry(20,40,400,40)
		
        lbl3 = QLabel(self)
        lbl3.setText('Make sure that your pc is connected to the internet')
        lbl3.setGeometry(20,60,400,40)
		
        # Create textbox
        self.textbox = QLineEdit(self) # Textbox Space object creation.
        self.textbox.move(20, 100) # set the position of the textbox in prompt screen.
        self.textbox.resize(280,40) # set the size of the textbox.
 
        # Create a button in the window
        self.button = QPushButton('Make Article', self) # Push Button object creation.
        self.button.move(20,160) # set the position of the button on prompt screen.
 
        # connect button to function on_click
        self.button.clicked.connect(self.on_click) # set the Event listener.
        self.show()
		
    def shw_wt(self):
        # for the showing of label Please wait
        lbl4 = QLabel(self)
        lbl4.setText('Please wait')
        lbl4.setGeometry(20,135,400,40)
        lbl4.show()    


    # on_click: On click of the button the event handler brings the program to this fucntion. 
    @pyqtSlot()
    def on_click(self):
        """for the showing of label Please wait
        lbl4 = QLabel(self)
        lbl4.show()
        lbl4.setText('Please wait')
        lbl4.setGeometry(20,135,400,40)"""
        self.shw_wt()
        cur_topic = self.textbox.text() # Retreive the topic entered in the text box.
        topics = cur_topic.split(",") # Storing each topic part separated by ',' in a list.
        for ech_topic in topics: # searching each part of topic in wikipedia
            try:
                para = wikipedia.summary(ech_topic, sentences = 100) # Retreiving only 100 sentences of the topic from wiki.
            except wikipedia.exceptions.DisambiguationError as e:
                para = wikipedia.summary(e.options[0], sentences = 100)
            except requests.exceptions.ConnectionError as e:
                pop_up = QMessageBox.question(self, 'Warning!', 'You are not connected to the internet', QMessageBox.Ok)
                if pop_up == QMessageBox.Ok:
                   self.initUI()
                   # problem in above line
                else:
                   pass
            doc.add_heading (ech_topic,0) # creating each topic docs.
            doc.add_paragraph(para) # Adding the retreived  information into the docs.
        n = str(time.time())
        doc.save(n+".docx") # saving the docs with current time and .docx extension.
        pop_up = QMessageBox.question(self, 'All Done', 'Your article is made. Check your current directory', QMessageBox.Ok) # prompt message to display the success message.
        if pop_up == QMessageBox.Ok:
            pass
        else:
            pass
        #lbl4.hide()
        self.show()
		
 
# Main  fucntion: Calling of the App class.
if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = App()
    sys.exit(app.exec_())