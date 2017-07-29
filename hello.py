import sys
import time
import wikipedia
import docx
import queue
import requests
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtCore import *

doc = docx.Document()

class App(QMainWindow):

    def start(self):
	    #hello
        print("")
 
    def __init__(self):
        super().__init__()
        self.title = 'Articulate - tecbeast.com'
        self.left = 50
        self.top = 100
        self.width = 400
        self.height = 300
        self.initUI()
 
    def initUI(self):
        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)
        lbl1 = QLabel(self)
        lbl1.setText("Please enter the topics you want to print articles on")
        lbl1.setGeometry(20,20,400,40)

        lbl2 = QLabel(self)
        lbl2.setText('Topics should be separeted by comma ","')
        lbl2.setGeometry(20,40,400,40)
		
        lbl3 = QLabel(self)
        lbl3.setText('Make sure that your pc is connected to the internet')
        lbl3.setGeometry(20,60,400,40)
		
        # Create textbox
        self.textbox = QLineEdit(self)
        self.textbox.move(20, 100)
        self.textbox.resize(280,40)
 
        # Create a button in the window
        self.button = QPushButton('Make Article', self)
        self.button.move(20,160)
 
        # connect button to function on_click
        self.button.clicked.connect(self.on_click)
        self.show()
		
    def shw_wt(self):
        #for the showing of label Please wait
        lbl4 = QLabel(self)
        lbl4.setText('Please wait')
        lbl4.setGeometry(20,135,400,40)
        lbl4.show()    


 
    @pyqtSlot()
    def on_click(self):
        """for the showing of label Please wait
        lbl4 = QLabel(self)
        lbl4.show()
        lbl4.setText('Please wait')
        lbl4.setGeometry(20,135,400,40)"""
        self.shw_wt()
        cur_topic = self.textbox.text()
        topics = cur_topic.split(",")
        for ech_topic in topics:
            try:
                para = wikipedia.summary(ech_topic, sentences = 100)
            except wikipedia.exceptions.DisambiguationError as e:
                para = wikipedia.summary(e.options[0], sentences = 100)
            except requests.exceptions.ConnectionError as e:
                pop_up = QMessageBox.question(self, 'Warning!', 'You are not connected to the internet', QMessageBox.Ok)
                if pop_up == QMessageBox.Ok:
                   self.initUI()
                   #problem in above line
                else:
                   pass
            doc.add_heading (ech_topic,0)
            doc.add_paragraph(para)
        n = str(time.time())
        doc.save(n+".docx")
        pop_up = QMessageBox.question(self, 'All Done', 'Your article is made. Check your current directory', QMessageBox.Ok)
        if pop_up == QMessageBox.Ok:
            pass
        else:
            pass
        #lbl4.hide()
        self.show()
		
 
 
if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = App()
    sys.exit(app.exec_())